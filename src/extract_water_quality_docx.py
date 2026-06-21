"""Extract water-quality monitoring tables from the supplied DOCX appendix."""

from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path

import docx


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = r"F:\心理学接单\sci ssci\推荐过来的三区\2、附表报告(监测原始数据）.docx"


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\n", " ")).strip()


def parse_date_from_title(title: str) -> str:
    match = re.search(r"(20\d{2})年(\d{1,2})月", title)
    if not match:
        return ""
    return f"{match.group(1)}-{int(match.group(2)):02d}"


def parse_river_from_title(title: str) -> str:
    stripped = re.sub(r"^附表\s*\d+\s*", "", title)
    stripped = re.sub(r"20\d{2}年\d{1,2}月.*$", "", stripped).strip()
    return stripped


def extract_titles(document: docx.Document) -> list[str]:
    titles = []
    for para in document.paragraphs:
        text = clean_text(para.text)
        if re.match(r"^附表\s*\d+\s*", text) and "水质检测结果" in text:
            titles.append(text)
    if len(titles) >= 66:
        return titles[-33:]
    return titles[-33:] if len(titles) >= 33 else titles


def extract_rows(docx_path: Path) -> list[dict[str, str]]:
    document = docx.Document(docx_path)
    titles = extract_titles(document)
    rows = []

    for table_index, table in enumerate(document.tables):
        title = titles[table_index] if table_index < len(titles) else f"附表 {table_index + 1}"
        date = parse_date_from_title(title)
        river = parse_river_from_title(title)
        if len(table.rows) < 3 or len(table.columns) < 3:
            continue

        report_numbers = [clean_text(cell.text) for cell in table.rows[0].cells]
        station_names = [clean_text(cell.text) for cell in table.rows[1].cells]

        for col_idx in range(2, len(table.columns)):
            report_number = report_numbers[col_idx] if col_idx < len(report_numbers) else ""
            station_name = station_names[col_idx] if col_idx < len(station_names) else ""
            for row_idx in range(2, len(table.rows)):
                cells = [clean_text(cell.text) for cell in table.rows[row_idx].cells]
                if len(cells) < 3:
                    continue
                indicator = cells[0]
                unit = cells[1]
                value = cells[col_idx] if col_idx < len(cells) else ""
                if not indicator or not station_name:
                    continue
                rows.append(
                    {
                        "appendix_table": str(table_index + 1),
                        "caption": title,
                        "sampling_month": date,
                        "river_or_group": river,
                        "report_number": report_number,
                        "station_name": station_name,
                        "indicator": indicator,
                        "unit": unit,
                        "value_raw": value,
                    }
                )

    return rows


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default=DEFAULT_DOCX)
    parser.add_argument("--output", default="data/raw/water_quality_monitoring_long.csv")
    args = parser.parse_args()

    rows = extract_rows(Path(args.input))
    output = ROOT / args.output
    output.parent.mkdir(parents=True, exist_ok=True)

    fields = [
        "appendix_table",
        "caption",
        "sampling_month",
        "river_or_group",
        "report_number",
        "station_name",
        "indicator",
        "unit",
        "value_raw",
    ]
    with output.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)

    print(f"Wrote {len(rows)} rows to {output}")


if __name__ == "__main__":
    main()
